import streamlit as st
import streamlit.components.v1 as components
from PIL import Image, ImageDraw, ImageOps
import io
from docx import Document
from docx.shared import Inches
from streamlit_drawable_canvas import st_canvas

# 用于保存截图和批注的全局变量
if "snapshots" not in st.session_state:
    st.session_state.snapshots = []
if "video_url" not in st.session_state:
    st.session_state.video_url = None
if "uploaded_image" not in st.session_state:
    st.session_state.uploaded_image = None
if "edited_image" not in st.session_state:
    st.session_state.edited_image = None

def add_snapshot(image_data, video_time, annotation):
    st.session_state.snapshots.append({"image": image_data, "time": video_time, "annotation": annotation})

def delete_snapshot(index):
    del st.session_state.snapshots[index]

def update_annotation(index, new_annotation):
    st.session_state.snapshots[index]["annotation"] = new_annotation

def export_to_docx():
    doc = Document()
    doc.add_heading('Video Notes', 0)
    for snapshot in st.session_state.snapshots:
        doc.add_picture(io.BytesIO(snapshot["image"]), width=Inches(4))
        doc.add_paragraph(f"Time: {snapshot['time']} seconds")
        doc.add_paragraph(snapshot["annotation"])
        doc.add_paragraph("\n")
    return doc

def main():
    st.set_page_config(layout="wide")
    st.title("AI视频总结 WebApp")

    left_column, right_column = st.columns([3, 2])

    with left_column:
        url = st.text_input("输入 YouTube 视频 URL")
        generate_button = st.button("生成")

        if generate_button and url:
            video_id = url.split('v=')[1]
            st.session_state.video_url = f"https://www.youtube.com/embed/{video_id}?enablejsapi=1"

        if st.session_state.video_url:
            components.html(f"""
                <iframe id="videoFrame" width="100%" height="600" src="{st.session_state.video_url}" frameborder="0" allow="autoplay; encrypted-media" allowfullscreen></iframe>
            """, height=600)

    with right_column:
        st.markdown("## 做笔记")

        # 上传手动截图
        uploaded_file = st.file_uploader("上传手动截图", type=["png", "jpg", "jpeg"])
        if uploaded_file is not None:
            st.session_state.uploaded_image = uploaded_file.read()

        if st.session_state.uploaded_image:
            img = Image.open(io.BytesIO(st.session_state.uploaded_image)).convert("RGBA")
            st.image(st.session_state.uploaded_image, caption="Uploaded Frame", use_column_width=True)
            
            # 编辑图片
            drawing_mode = st.selectbox("选择绘制模式", ("freedraw", "line", "rect", "circle", "transform"))
            stroke_width = st.slider("笔刷粗细: ", 1, 25, 3)
            color = st.color_picker("选择颜色: ", "#000000")

            canvas_result = st_canvas(
                fill_color="rgba(255, 165, 0, 0.3)",  # 背景填充颜色
                stroke_width=stroke_width,
                stroke_color=color,
                background_image=img,
                update_streamlit=True,
                height=480,
                drawing_mode=drawing_mode,
                key="canvas",
            )

            # 保存绘制后的图片
            if st.button("保存编辑后的图片"):
                if canvas_result.image_data is not None:
                    # Convert the canvas drawing to an image
                    edited_image = Image.fromarray(canvas_result.image_data.astype("uint8"), "RGBA")
                    
                    # Ensure both images are the same size
                    if img.size != edited_image.size:
                        edited_image = edited_image.resize(img.size)
                    
                    # Create a new image for combining
                    combined_image = Image.alpha_composite(img, edited_image)
                    
                    # Save the result
                    img_byte_arr = io.BytesIO()
                    combined_image.save(img_byte_arr, format="PNG")
                    st.session_state.edited_image = img_byte_arr.getvalue()
                    st.success("图片已保存")
                else:
                    st.warning("请在画布上绘制一些内容后再保存")

        # 输入视频时间和批注
        video_time = st.number_input("输入视频时间（秒）", min_value=0)
        annotation = st.text_area("输入批注内容")
        add_note_button = st.button("添加批注")

        if add_note_button and st.session_state.edited_image and video_time is not None and annotation:
            add_snapshot(st.session_state.edited_image, video_time, annotation)
            st.success(f"已添加批注: {video_time}秒 - {annotation}")

        st.markdown("### 已添加的批注")

        # 使用容器来使批注部分可滚动
        with st.container():
            for i, snapshot in enumerate(st.session_state.snapshots):
                with st.expander(f"批注 {i+1}: {snapshot['time']} 秒"):
                    st.image(snapshot["image"], caption=f"{snapshot['time']}秒 - {snapshot['annotation']}")
                    new_annotation = st.text_area(f"编辑批注 {i+1}", value=snapshot["annotation"], key=f"edit_{i}")
                    if st.button("更新批注", key=f"update_{i}"):
                        update_annotation(i, new_annotation)
                        st.success("批注已更新")
                    
                    # 删除批注后更新页面
                    if st.button("删除", key=f"delete_{i}"):
                        delete_snapshot(i)
                        st.experimental_rerun()  # 替换为 st._rerun() 来刷新页面

                    if st.button(f"跳转到 {snapshot['time']} 秒", key=f"jump_{i}"):
                        components.html(f"""
                            <script>
                                seekToTime({int(snapshot['time'])});
                            </script>
                        """, height=0)

        # 导出功能
        st.markdown("## 导出笔记")
        docx_export_button = st.button("导出为 DOCX")

        if docx_export_button:
            doc = export_to_docx()
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            st.download_button("下载 DOCX 文件", data=buffer, file_name="video_notes.docx")

if __name__ == "__main__":
    main()
